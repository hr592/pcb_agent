# IMPORTS
import json # reads components.json and turns it into a Python list
import os # to read environment variables like the API key
import io # creates in-memory file buffers for downloads (no saving to disk)
from groq import Groq # gives us the Groq class to create an API client
import streamlit as st # the entire UI framework
from dotenv import load_dotenv # reads .env file and loads GROQ_API_KEY into memory
import pandas as pd # build the CSV BOM export
from docx import Document # python-docx for Word documents
from docx.shared import Pt, RGBColor # font size and colour tools for Word docs
from docx.enum.text import WD_ALIGN_PARAGRAPH # text alignment options for Word docs

# SETUP
load_dotenv() # reads .env file

# creates one Groq client object reused for every API call in the app
# os.getenv("GROQ_API_KEY") reads the key from .env
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# PAGE CONFIG
st.set_page_config(page_title="PCB Component Agent", page_icon="🔧", layout="wide") # sets browser tab title, icon and wide layout
st.title("🔧 PCB Component Agent") # big heading at the top of the page
st.caption("Describe what you need and the agent will find the best match from your parts library.") # small subtitle text

# LOAD COMPONENTS
@st.cache_data # caches the result so this function only runs ONCE
def load_components():
    with open("components.json") as f: # opens the file
        return json.load(f) # parses the JSON and returns a Python list of dictionaries

components = load_components() # stores the full parts list in memory

# SESSION STATE INITIALIZATION
# session_state is a dictionary that persists between reruns
# The "if not in" guard means we only initialize each variable ONCE
if "approved_parts" not in st.session_state:
    st.session_state.approved_parts = [] # list of part dicts the user has approved

if "messages" not in st.session_state:
    st.session_state.messages = [] # conversation history sent to the AI on every call (user + assistant messages)

if "current_recommendation" not in st.session_state:
    st.session_state.current_recommendation = None # the AI's last response text

if "filtered_parts" not in st.session_state:
    st.session_state.filtered_parts = [] # the filtered parts list saved at search time, reused by the reject loop

# PRESETS
# Dictionary of preset configurations
PRESETS = {
    "Custom": None, # use whatever sliders set to
    "💰 Cost Optimized":   {"max_price": 0.50, "min_stock": 1000}, # change sliders to these settings if selected
    "🔒 High Reliability": {"max_price": 10.00, "min_stock": 500},
    "⚡ Prototype Fast":   {"max_price": 5.00,  "min_stock": 1},
}

if "preset" not in st.session_state:
    st.session_state.preset = "Custom"  # default preset on first load

# SIDEBAR
# everything inside this block is in the left sidebar (for python to filter before AI)
with st.sidebar:
    st.header("⚙️ Constraints")

    st.subheader("Quick Presets")
    preset_choice = st.radio( # radio buttons so user can pick one option
        "Start from a preset:",
        options=list(PRESETS.keys()),
        index=list(PRESETS.keys()).index(st.session_state.preset),  # keeps the current selection highlighted
        horizontal=False,
        label_visibility="collapsed" # hides the "Start from a preset:" label
    )
    if preset_choice != st.session_state.preset: # only update if the user actually changed it
        st.session_state.preset = preset_choice # save the new preset to session state
        st.rerun() # rerun the script so the sliders update to preset values

    st.divider()
    st.caption("Or customize manually:")

    active_preset = PRESETS[st.session_state.preset] # gets the preset dict (or None if Custom)

    # builds a sorted list of unique package sizes from components.json
    all_packages = sorted(set(p["supplier_device_package"] for p in components))
    selected_packages = st.multiselect( # user can select multiple options
        "Allowed package sizes",
        options=all_packages,
        default=all_packages, # all packages selected by default
    )

    # if a preset is active use its max_price, otherwise default to 5.00
    default_price = active_preset["max_price"] if active_preset else 5.00
    max_price = st.slider(
        "Max price per part (CAD)",
        min_value=0.10, max_value=10.00,
        value=default_price, # this is what makes the preset actually change the slider
        step=0.05, format="$%.2f"
    )

    # same pattern for min_stock
    default_stock = active_preset["min_stock"] if active_preset else 100
    min_stock = st.number_input(
        "Minimum stock quantity",
        min_value=1,
        value=default_stock,
        step=100
    )

    all_categories = sorted(set(p["category"] for p in components))  # unique categories from components.json
    selected_categories = st.multiselect(
        "Part categories", options=all_categories, default=all_categories
    )

    st.divider()

    # PYTHON PRE-FILTER
    # This filters the full parts list using the sidebar values BEFORE any AI call
    filtered_preview = [
        p for p in components
        if p["supplier_device_package"] in selected_packages # package must be in the allowed list
        and p["price"] <= max_price # price must be under the max
        and p["stock"] >= min_stock # stock must be above the minimum
        and p["category"] in selected_categories # category must be selected
    ]
    st.metric("Parts matching your constraints", len(filtered_preview)) # shows the count live
    if len(filtered_preview) == 0:
        st.warning("⚠️ No parts match. Please try relaxing your constraints.") # warns if nothing passes

    st.divider()
    if st.button("🔄 Start New Search"):
        # resets all session state so the user can start fresh
        st.session_state.current_recommendation = None
        st.session_state.messages = []
        st.session_state.filtered_parts = []
        st.session_state.preset = "Custom"
        st.session_state["reject_reason"] = "" # clears the reject reason text box
        st.rerun() # reruns the script with everything cleared

# HELPER FUNCTIONS

def build_constraint_summary(packages, max_p, min_s, categories):
    # turns the sidebar values into plain English that gets injected into the AI system prompt
    lines = [
        f"- Allowed packages: {', '.join(packages) if packages else 'any'}",
        f"- Max price per part: ${max_p:.2f}",
        f"- Minimum stock quantity: {min_s:,}",
        f"- Allowed categories: {', '.join(categories) if categories else 'any'}",
    ]
    return "\n".join(lines)

def call_ai(filtered_parts, constraint_summary):
    # builds and sends the API request to Groq for BOTH the initial search and the reject loop
    system_message = {
        "role": "system",
        # the system prompt
        "content": f"""You are an expert PCB component selection assistant.
Your job is to recommend the single best matching part from the list below.

The user has set these manufacturing constraints. You must respect them in your reasoning:
{constraint_summary}

Rules:
- Only recommend parts from the provided list — never invent or hallucinate parts
- Explain WHY the part fits the user's requirements AND constraints
- If this is a follow-up after a rejection, suggest a DIFFERENT part than previously rejected
- If the user gave a rejection reason, take it into account when picking the alternative
- Always end your response with a line that says exactly:
  RECOMMENDED_PART: <digikey_part_number>
- Be concise and technical

Available parts (already filtered by user constraints):
{json.dumps(filtered_parts, indent=2)}""" # sends the filtered parts list as formatted JSON text
    }
    # combine system message with the full conversation history
    all_messages = [system_message] + st.session_state.messages
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile", # the specific Groq model to use
        max_tokens=1024, # maximum length of the AI's response
        messages=all_messages # the full conversation
    )
    return response.choices[0].message.content # extracts just the text from the API response object

def extract_recommended_part(recommendation_text, parts_list):
    # parses the AI response to find which part was actually recommended
    for line in recommendation_text.split("\n"): # loops through each line of the response
        if "RECOMMENDED_PART:" in line: # finds the special tag line
            part_number = line.split("RECOMMENDED_PART:")[-1].strip() # extracts the part number after the colon
            for part in parts_list:
                if part["digikey_part_number"] == part_number: # finds the matching part in our JSON
                    return part # returns the full part dict
    return None

def generate_csv(approved_parts):
    # builds a CSV from the approved parts list using pandas
    rows = []
    for i, part in enumerate(approved_parts): # enumerate gives both index (i) and value (part)
        rows.append({
            "Designator":   f"{part['category'][0]}{i+1}",
            "Comment":      part.get("resistance") or part.get("capacitance") or part.get("function") or part["description"],
            "Footprint":    part["supplier_device_package"],
            "DigiKey PN":   part["digikey_part_number"],
            "MFR PN":       part["manufacturer_part_number"],
            "Manufacturer": part["manufacturer"],
            "Description":  part["description"],
            "Quantity":     1,
            "Unit Price":   part["price"],
            "Stock":        part["stock"],
        })
    df = pd.DataFrame(rows) # converts the list of dicts into a pandas DataFrame (like a spreadsheet in memory)
    return df.to_csv(index=False)

def generate_word_doc(approved_parts):
    # builds a Word document summary of all approved parts using python-docx
    doc = Document()
    title = doc.add_heading("PCB Component Summary", 0) # adds a title heading (level 0 = biggest)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Total approved parts: {len(approved_parts)}")
    doc.add_paragraph("")  # blank line for spacing

    for i, part in enumerate(approved_parts):
        doc.add_heading(f"{i+1}. {part['manufacturer_part_number']}", level=1) # section heading per part
        doc.add_paragraph(part["description"])

        table = doc.add_table(rows=1, cols=2) # creates a 2-column table starting with 1 row (the header)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"

        fields = [
            ("Category",       part.get("category")),
            ("Manufacturer",   part.get("manufacturer")),
            ("Package",        part.get("supplier_device_package")),
            ("Price",          f"${part.get('price', 0):.2f}"),
            ("Stock",          f"{part.get('stock', 0):,}"),
            ("Resistance",     part.get("resistance")),
            ("Capacitance",    part.get("capacitance")),
            ("Tolerance",      part.get("tolerance")),
            ("Voltage Rating", part.get("voltage_rating")),
            ("Power Rating",   part.get("power_rating")),
            ("Function",       part.get("function")),
            ("Temp Range",     part.get("operating_temperature")),
        ]
        for label, value in fields:
            if value: # only adds a row if the field has data
                row = table.add_row().cells
                row[0].text = label
                row[1].text = str(value) # str() converts numbers to strings for the table

        doc.add_paragraph("")
        doc.add_paragraph(f"DigiKey: {part.get('digikey_link', 'N/A')}")
        doc.add_paragraph(f"Datasheet: {part.get('datasheet_link', 'N/A')}")
        doc.add_paragraph("")

    # saves the document to an in-memory buffer instead of a file on disk
    # io.BytesIO() acts like a file but lives in RAM
    buffer = io.BytesIO()
    doc.save(buffer) # writes the Word doc into the buffer
    buffer.seek(0)
    return buffer

# MAIN AREA
st.subheader("🔍 Describe what you need")

user_query = st.text_input(
    "What component are you looking for?",
    placeholder="e.g. a decoupling capacitor for a 3.3V rail, 0402 or smaller",
)

search_button = st.button(
    "🔎 Find Best Part", type="primary",
    disabled=len(filtered_preview) == 0  # button is greyed out if no parts pass the current constraints
)

constraint_summary = build_constraint_summary(
    selected_packages, max_price, min_stock, selected_categories
)

# INITIAL SEARCH
# this block only runs when the user clicks "Find Best Part" AND has typed something
if search_button and user_query:
    st.session_state.messages = [] # clears old conversation history for a fresh search
    st.session_state.current_recommendation = None # clears any previous recommendation from the screen
    st.session_state.filtered_parts = filtered_preview # saves the filtered parts to session state for the reject loop to use later
    st.session_state.constraint_summary = constraint_summary # saves constraints for the reject loop too

    # adds the user's input to the conversation history as the first message
    st.session_state.messages.append({"role": "user", "content": user_query})

    with st.spinner("🤖 AI is finding the best part..."): # shows a loading spinner while waiting for the API
        try:
            recommendation_text = call_ai( # calls the AI
                st.session_state.filtered_parts,
                st.session_state.constraint_summary
            )
            st.session_state.messages.append({"role": "assistant", "content": recommendation_text}) # saves AI response to history
            st.session_state.current_recommendation = recommendation_text
        except Exception as e:
            st.error(f"API error: {e}")
            st.stop()

# DISPLAY RECOMMENDATION
# this block runs on every rerun as long as there's a recommendation in session state
if st.session_state.current_recommendation:
    st.divider()
    st.subheader("🤖 AI Recommendation")
    st.markdown(st.session_state.current_recommendation)

    part = extract_recommended_part(
        st.session_state.current_recommendation,
        st.session_state.filtered_parts
    )
    if part:
        # shows a quick part card with key info so the user can review before deciding
        st.divider()
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Price", f"${part['price']:.2f}")
        with col2:
            st.metric("Stock", f"{part['stock']:,}")
        with col3:
            st.metric("Package", part['supplier_device_package'])
        st.markdown(f"**Manufacturer:** {part['manufacturer']} &nbsp;|&nbsp; **MFR PN:** {part['manufacturer_part_number']}")
        st.markdown(f"[🛒 View on DigiKey]({part['digikey_link']}) &nbsp;&nbsp; [📄 Datasheet]({part['datasheet_link']})")  # clickable links
        st.divider()

    col1, col2 = st.columns(2)

    # APPROVE
    with col1:
        if st.button("✅ Approve this part", type="primary"):
            if part:
                # checks if this part is already in the approved list to prevent duplicates
                already_approved = any(
                    p["digikey_part_number"] == part["digikey_part_number"]
                    for p in st.session_state.approved_parts
                )
                if not already_approved:
                    st.session_state.approved_parts.append(part) # adds the full part dict to the BOM list
                    st.success(f"✅ Added **{part['manufacturer_part_number']}** to your BOM!")
                else:
                    st.warning("This part is already in your approved list.")
            else:
                st.error("Couldn't identify the part — try searching again.")
            # clears the recommendation and conversation so the user can search for the next part
            st.session_state.current_recommendation = None
            st.session_state.messages = []
            st.rerun()

    # REJECT WITH REASON
    with col2:
        reject_reason = st.text_input(
            "Reason for rejection (optional):",
            placeholder="e.g. too expensive, need higher voltage rating, prefer KEMET",
            key="reject_reason"
        )
        if st.button("❌ Reject — find alternative"):
            # builds the rejection message
            if reject_reason:
                rejection_msg = f"I reject this part. Reason: {reject_reason}. Please suggest a different part that addresses this concern."
            else:
                rejection_msg = "I reject this part. Please suggest a different one from the list."

            # adds the rejection to conversation history
            st.session_state.messages.append({"role": "user", "content": rejection_msg})

            with st.spinner("🤖 Finding an alternative..."):
                try:
                    new_recommendation = call_ai(
                        st.session_state.filtered_parts,
                        st.session_state.get("constraint_summary", constraint_summary)
                    )
                    st.session_state.messages.append({"role": "assistant", "content": new_recommendation})
                    st.session_state.current_recommendation = new_recommendation
                    st.rerun()
                except Exception as e:
                    st.error(f"API error: {e}")

# APPROVED PARTS AND EXPORT
if st.session_state.approved_parts: # only shows this section if at least one part has been approved
    st.divider()
    st.subheader(f"✅ Approved Parts ({len(st.session_state.approved_parts)})")

    for i, part in enumerate(st.session_state.approved_parts):
        with st.expander(f"**{part['manufacturer_part_number']}** — {part['description']}"): # collapsible section per part
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Price", f"${part['price']:.2f}")
            with col2:
                st.metric("Stock", f"{part['stock']:,}")
            with col3:
                st.metric("Package", part['supplier_device_package'])
            st.markdown(f"[📄 Datasheet]({part['datasheet_link']}) &nbsp;&nbsp; [🛒 DigiKey]({part['digikey_link']})")

    st.divider()
    st.subheader("📦 Export BOM")
    col1, col2 = st.columns(2)
    with col1:
        csv_data = generate_csv(st.session_state.approved_parts) # generates the CSV string
        st.download_button(
            label="⬇️ Download CSV BOM",
            data=csv_data,
            file_name="pcb_bom.csv",
            mime="text/csv",
        )
    with col2:
        word_buffer = generate_word_doc(st.session_state.approved_parts) # generates the Word doc in memory
        st.download_button(
            label="⬇️ Download Word Summary",
            data=word_buffer,
            file_name="pcb_component_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
else:
    st.divider()
    st.caption("Approved parts will appear here. Export to BOM when you're done.")